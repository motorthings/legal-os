from logger_config import get_logger
logger = get_logger(__name__)

"""
Google Drive Sync Service

Handles syncing documents from Google Drive to SuperAssistant using LlamaIndex.
Manages OAuth tokens, document fetching, and integration with existing document processor.

Flow:
1. Get/refresh OAuth token from database
2. Initialize LlamaIndex GoogleDriveReader
3. Fetch documents from specified folder (or entire Drive)
4. Create/update document records in database
5. Save document content to Supabase storage
6. Trigger existing document processor for embedding generation
"""

import os
import io
import re
import time
import signal
import threading
from typing import Optional, Dict, List, Any
from datetime import datetime, timedelta, timezone
import uuid
from uuid import UUID
from contextlib import contextmanager

from dotenv import load_dotenv
from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from llama_index.readers.google import GoogleDriveReader

# Simple .docx reader using python-docx (already installed)
from docx import Document as DocxDocument
from llama_index.core.readers.base import BaseReader
from llama_index.core.schema import Document as LlamaDocument
from pathlib import Path
from typing import List

class SimpleDocxReader(BaseReader):
    """Simple .docx reader using python-docx library."""

    def load_data(self, file: Path, extra_info: dict = None) -> List[LlamaDocument]:
        """Load data from .docx file."""
        doc = DocxDocument(file)
        text = '\n\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
        return [LlamaDocument(text=text, metadata=extra_info or {})]

from database import get_supabase
from services.oauth_crypto import encrypt_token, decrypt_token, OAuthCryptoError
from document_processor import process_document
from config import get_client_id_for_user

load_dotenv()

# Get Supabase configuration
SUPABASE_URL = os.getenv("SUPABASE_URL", "")

# Get centralized Supabase client
supabase = get_supabase()

# Pre-compiled regex patterns for text validation (performance optimization)
_XML_PATTERN = re.compile(r'<[^>]+>')
_WORD_PATTERN = re.compile(r'[a-zA-Z]{2,}')
_SUSPICIOUS_PATTERNS = [
    (re.compile(r'word/document\.xml'), '.docx document.xml'),
    (re.compile(r'word/numbering\.xml'), '.docx numbering.xml'),
    (re.compile(r'word/styles\.xml'), '.docx styles.xml'),
]


class GoogleDriveSyncError(Exception):
    """Raised when Google Drive sync operations fail"""
    pass


class TimeoutError(Exception):
    """Raised when an operation times out"""
    pass


@contextmanager
def timeout(seconds):
    """
    Context manager to add timeout to operations.

    Note: This uses signal.SIGALRM which only works on Unix systems and
    only in the main thread. When called from APScheduler background threads,
    this will be a no-op and won't enforce timeouts.

    Args:
        seconds: Maximum time to allow operation to run

    Raises:
        TimeoutError: If operation exceeds timeout
    """
    def timeout_handler(signum, frame):
        raise TimeoutError(f"Operation timed out after {seconds} seconds")

    # Only set timeout if we're on a Unix system and in the main thread
    # Otherwise, just run the code without timeout protection
    try:
        if hasattr(signal, 'SIGALRM') and threading.current_thread() is threading.main_thread():
            old_handler = signal.signal(signal.SIGALRM, timeout_handler)
            signal.alarm(seconds)
            timeout_enabled = True
        else:
            timeout_enabled = False

        yield

    finally:
        if timeout_enabled:
            signal.alarm(0)
            signal.signal(signal.SIGALRM, old_handler)



# ============================================================================
# Token Management
# ============================================================================

def get_user_tokens(user_id: str) -> Optional[Dict]:
    """
    Get OAuth tokens for a user from database.

    Args:
        user_id: UUID of the user

    Returns:
        Dict with token info or None if not connected

    Raises:
        GoogleDriveSyncError: If token retrieval fails
    """
    try:
        result = supabase.table('google_drive_tokens')\
            .select('*')\
            .eq('user_id', user_id)\
            .eq('is_active', True)\
            .execute()

        if not result.data:
            return None

        return result.data[0]

    except Exception as e:
        raise GoogleDriveSyncError(f"Failed to get user tokens: {e}")


def get_or_refresh_token(user_id: str) -> Credentials:
    """
    Get valid OAuth credentials, refreshing if necessary.

    Args:
        user_id: UUID of the user

    Returns:
        google.oauth2.credentials.Credentials object

    Raises:
        GoogleDriveSyncError: If no tokens found or refresh fails
    """
    token_record = get_user_tokens(user_id)

    if not token_record:
        raise GoogleDriveSyncError(
            "No Google Drive connection found. Please connect your Google Drive first."
        )

    try:
        # Decrypt tokens
        access_token = decrypt_token(token_record['access_token_encrypted'])
        refresh_token = decrypt_token(token_record['refresh_token_encrypted'])

        # Create credentials object
        credentials = Credentials(
            token=access_token,
            refresh_token=refresh_token,
            token_uri="https://oauth2.googleapis.com/token",
            client_id=os.environ.get("GOOGLE_CLIENT_ID"),
            client_secret=os.environ.get("GOOGLE_CLIENT_SECRET"),
            scopes=token_record.get('scopes', [
                'https://www.googleapis.com/auth/drive.readonly',
                'https://www.googleapis.com/auth/drive.metadata.readonly'
            ])
        )

        # Check if token is expired or about to expire (within 5 minutes)
        token_expires_at = datetime.fromisoformat(
            token_record['token_expires_at'].replace('Z', '+00:00')
        )
        now = datetime.now(timezone.utc)

        if now >= token_expires_at - timedelta(minutes=5):
            logger.info(f"   🔄 Token expired, refreshing...")

            # Refresh the token
            credentials.refresh(Request())

            # Update database with new tokens
            update_data = {
                'access_token_encrypted': encrypt_token(credentials.token),
                'token_expires_at': (now + timedelta(seconds=3600)).isoformat(),
                'last_refreshed_at': now.isoformat()
            }

            # If we got a new refresh token, update it too
            if credentials.refresh_token:
                update_data['refresh_token_encrypted'] = encrypt_token(
                    credentials.refresh_token
                )

            supabase.table('google_drive_tokens')\
                .update(update_data)\
                .eq('user_id', user_id)\
                .execute()

            logger.info(f"   ✓ Token refreshed successfully")

        return credentials

    except OAuthCryptoError as e:
        raise GoogleDriveSyncError(f"Token decryption failed: {e}")
    except Exception as e:
        raise GoogleDriveSyncError(f"Token refresh failed: {e}")


# ============================================================================
# Text Validation
# ============================================================================

def is_valid_text_content(text: str, min_printable_ratio: float = 0.70) -> bool:
    """
    Validate that text content is actually readable text, not binary garbage or corrupted .docx XML.

    Enhanced validation catches:
    - ZIP file headers (corrupted .docx exports)
    - Excessive XML tags (.docx internals)
    - Low word density (gibberish vs natural language)
    - Binary patterns that appear printable

    Args:
        text: Text to validate
        min_printable_ratio: Minimum ratio of printable characters (default 70%)

    Returns:
        True if text appears to be valid, False if binary/corrupted
    """
    # Check 1: Empty or whitespace only
    if not text or len(text.strip()) == 0:
        return False

    # Check 2: ZIP file header (corrupted .docx)
    if text.startswith('PK'):
        logger.warning(f"      ⚠️  ZIP file header detected (corrupted .docx export)")
        return False

    # Check 3: Printable ratio (catch obvious binary)
    printable_count = sum(1 for c in text if c.isprintable() or c.isspace())
    printable_ratio = printable_count / len(text) if len(text) > 0 else 0

    if printable_ratio < min_printable_ratio:
        logger.warning(f"      ⚠️  Low printable ratio: {printable_ratio:.2%} (threshold: {min_printable_ratio:.2%})")
        return False

    # Check 4: Excessive XML tags (.docx XML internals)
    # Count characters that are part of XML tags (using pre-compiled pattern)
    xml_chars = sum(len(match.group()) for match in _XML_PATTERN.finditer(text))
    xml_ratio = xml_chars / len(text) if len(text) > 0 else 0

    if xml_ratio > 0.30:  # More than 30% XML tags is suspicious
        logger.warning(f"      ⚠️  Excessive XML tags ({xml_ratio:.2%}) - likely .docx internals")
        return False

    # Check 5: Word density (natural language check)
    # Count sequences of 2+ alphabetic characters (words) using pre-compiled pattern
    words = _WORD_PATTERN.findall(text)
    word_chars = sum(len(word) for word in words)
    word_density = word_chars / len(text) if len(text) > 0 else 0

    if word_density < 0.10:  # Less than 10% alphabetic words is suspicious
        logger.warning(f"      ⚠️  Low word density ({word_density:.2%}) - not natural language")
        return False

    # Check 6: Suspicious patterns (common in corrupted files) using pre-compiled patterns
    for pattern, name in _SUSPICIOUS_PATTERNS:
        if pattern.search(text):
            logger.warning(f"      ⚠️  Suspicious pattern detected: {name}")
            return False

    return True


# ============================================================================
# Document Management
# ============================================================================

def get_document_by_google_id(user_id: str, google_drive_file_id: str) -> Optional[Dict]:
    """
    Check if a document with given Google Drive file ID already exists.

    Args:
        user_id: UUID of the user
        google_drive_file_id: Google Drive file ID

    Returns:
        Document record or None if not found
    """
    try:
        result = supabase.table('documents')\
            .select('*')\
            .eq('user_id', user_id)\
            .eq('google_drive_file_id', google_drive_file_id)\
            .execute()

        return result.data[0] if result.data else None

    except Exception as e:
        logger.warning(f"   ⚠️  Error checking for existing document: {e}")
        return None


def create_drive_document(
    user_id: str,
    client_id: str,
    filename: str,
    file_content: bytes,
    google_drive_file_id: str,
    external_url: str,
    mime_type: str = "text/plain"
) -> Dict:
    """
    Create a new document record and upload content to storage.

    Args:
        user_id: UUID of the user
        client_id: UUID of the client/organization
        filename: Name of the file
        file_content: File content as bytes
        google_drive_file_id: Google Drive file ID
        external_url: URL to view file in Google Drive
        mime_type: MIME type of the file

    Returns:
        Created document record

    Raises:
        GoogleDriveSyncError: If document creation fails
    """
    try:
        # Generate storage path with UUID to prevent collisions
        unique_id = str(uuid.uuid4())
        storage_path = f"{user_id}/{unique_id}_{filename}"

        # Upload to Supabase storage
        logger.info(f"      📤 Uploading to storage: {storage_path}")
        supabase.storage.from_('documents').upload(
            path=storage_path,
            file=file_content,
            file_options={"content-type": mime_type}
        )

        # Get public URL (or signed URL depending on your bucket settings)
        storage_url = f"{SUPABASE_URL}/storage/v1/object/public/documents/{storage_path}"

        # Create document record
        now = datetime.now(timezone.utc).isoformat()
        document_data = {
            'user_id': user_id,
            'client_id': client_id,
            'uploaded_by': user_id,  # Required for documents list filtering
            'filename': filename,
            'storage_url': storage_url,
            'source_platform': 'google_drive',
            'external_id': google_drive_file_id,
            'external_url': external_url,
            'google_drive_file_id': google_drive_file_id,
            'last_synced_at': now,
            'processed': False,
            'uploaded_at': now
        }

        result = supabase.table('documents').insert(document_data).execute()
        document = result.data[0]

        logger.info(f"      ✓ Created document record: {document['id']}")
        logger.debug(f"      Document data: uploaded_by={document.get('uploaded_by')}, client_id={document.get('client_id')}, user_id={document.get('user_id')}")
        return document

    except Exception as e:
        logger.error(f"      ❌ Failed to create document: {e}")
        logger.error(f"      Document data: {document_data}")
        raise GoogleDriveSyncError(f"Failed to create document: {e}")


def update_drive_document(
    document_id: str,
    file_content: bytes,
    storage_url: str
) -> Dict:
    """
    Update an existing document's content.

    Args:
        document_id: UUID of the document
        file_content: New file content as bytes
        storage_url: Current storage URL

    Returns:
        Updated document record

    Raises:
        GoogleDriveSyncError: If update fails
    """
    try:
        # Extract storage path from URL
        storage_path = storage_url.split('/documents/')[-1]

        # Update file in storage
        logger.info(f"      📤 Updating storage: {storage_path}")
        supabase.storage.from_('documents').update(
            path=storage_path,
            file=file_content,
            file_options={"upsert": "true"}
        )

        # Delete old embeddings
        logger.info(f"      🗑️  Deleting old embeddings...")
        supabase.table('document_chunks').delete().eq('document_id', document_id).execute()

        # Update document record
        now = datetime.now(timezone.utc).isoformat()
        update_data = {
            'last_synced_at': now,
            'processed': False,  # Will be re-processed
            'processed_at': None
        }

        result = supabase.table('documents')\
            .update(update_data)\
            .eq('id', document_id)\
            .execute()

        logger.info(f"      ✓ Updated document record")
        return result.data[0]

    except Exception as e:
        raise GoogleDriveSyncError(f"Failed to update document: {e}")


# ============================================================================
# File Discovery
# ============================================================================

def list_folder_files(user_id: str, folder_id: str) -> List[Dict]:
    """
    List all files in a Google Drive folder (without syncing them).

    Args:
        user_id: UUID of the user
        folder_id: Google Drive folder ID

    Returns:
        List of file metadata dicts with id, name, mimeType, webViewLink, size

    Raises:
        GoogleDriveSyncError: If listing fails
    """
    try:
        logger.info(f"\n📂 Listing files in folder for user {user_id}")
        logger.info(f"   Folder ID: {folder_id}")

        # Get valid credentials
        credentials = get_or_refresh_token(user_id)

        # Build Drive API service
        service = build('drive', 'v3', credentials=credentials)

        # List files in the folder
        query = f"'{folder_id}' in parents and trashed = false"
        results = service.files().list(
            q=query,
            pageSize=100,
            fields="files(id, name, mimeType, webViewLink, size, iconLink, thumbnailLink)"
        ).execute()

        files = results.get('files', [])

        # Filter out folders, only return files
        file_list = []
        for file in files:
            # Skip folders
            if file.get('mimeType') == 'application/vnd.google-apps.folder':
                continue

            file_list.append({
                'id': file.get('id'),
                'name': file.get('name'),
                'mimeType': file.get('mimeType'),
                'webViewLink': file.get('webViewLink'),
                'size': file.get('size'),
                'iconLink': file.get('iconLink'),
                'thumbnailLink': file.get('thumbnailLink')
            })

        logger.info(f"   ✓ Found {len(file_list)} files (excluded {len(files) - len(file_list)} folders)")
        return file_list

    except Exception as e:
        logger.error(f"   ❌ Failed to list folder files: {e}")
        raise GoogleDriveSyncError(f"Failed to list folder files: {e}")


# ============================================================================
# Sync Logic
# ============================================================================

def create_placeholder_documents(
    user_id: str,
    file_ids: List[str]
) -> List[Dict]:
    """
    Create placeholder document records immediately for selected files.
    This allows documents to appear in the UI with "Processing" status
    before the full content is downloaded.

    Args:
        user_id: UUID of the user
        file_ids: List of Google Drive file IDs

    Returns:
        List of created document records

    Raises:
        GoogleDriveSyncError: If creation fails
    """
    try:
        logger.info(f"\n📝 Creating placeholder documents for user {user_id}")
        logger.info(f"   File IDs: {len(file_ids)}")

        # Get valid credentials
        credentials = get_or_refresh_token(user_id)

        # Get user's client_id
        user_result = supabase.table('users').select('client_id').eq('id', user_id).execute()
        if not user_result.data:
            raise GoogleDriveSyncError(f"User {user_id} not found")

        client_id = user_result.data[0].get('client_id')

        # Build Drive API client to fetch metadata
        service = build('drive', 'v3', credentials=credentials)

        # Fetch existing documents to avoid duplicates
        existing_docs_result = supabase.table('documents')\
            .select('*')\
            .eq('user_id', user_id)\
            .in_('google_drive_file_id', file_ids)\
            .execute()

        existing_docs_map = {
            doc['google_drive_file_id']: doc
            for doc in existing_docs_result.data
            if doc.get('google_drive_file_id')
        }
        logger.info(f"   📋 Found {len(existing_docs_map)} existing documents")

        created_docs = []
        now = datetime.now(timezone.utc).isoformat()

        for file_id in file_ids:
            try:
                # Check if document already exists
                if file_id in existing_docs_map:
                    existing_doc = existing_docs_map[file_id]
                    logger.info(f"   ✓ Document already exists: {existing_doc['filename']} (ID: {existing_doc['id']})")
                    # Update last_synced_at and reset processed status for re-sync
                    supabase.table('documents').update({
                        'last_synced_at': now,
                        'processed': False,  # Mark as unprocessed for re-sync
                        'storage_url': f'placeholder/{file_id}'  # Reset to placeholder URL so sync task updates instead of creating new
                    }).eq('id', existing_doc['id']).execute()
                    created_docs.append(existing_doc)
                    continue

                # Fetch file metadata from Google Drive
                file_metadata = service.files().get(
                    fileId=file_id,
                    fields='id,name,mimeType,webViewLink'
                ).execute()

                filename = file_metadata.get('name', 'Untitled')
                mime_type = file_metadata.get('mimeType', 'text/plain')
                web_view_link = file_metadata.get('webViewLink', '')

                # Skip folders
                if mime_type == 'application/vnd.google-apps.folder':
                    logger.info(f"   📁 Skipping folder: {filename}")
                    continue

                # Create placeholder document record
                document_data = {
                    'user_id': user_id,
                    'client_id': client_id,
                    'uploaded_by': user_id,
                    'filename': filename,
                    'storage_url': f"placeholder/{file_id}",  # Temporary, will be updated
                    'source_platform': 'google_drive',
                    'external_id': file_id,
                    'external_url': web_view_link,
                    'google_drive_file_id': file_id,
                    'last_synced_at': now,
                    'processed': False,  # Will be set to True after processing
                    'uploaded_at': now
                }

                result = supabase.table('documents').insert(document_data).execute()
                document = result.data[0]

                created_docs.append(document)
                logger.info(f"   ✓ Created placeholder for: {filename} (ID: {document['id']})")

            except Exception as e:
                logger.error(f"   ❌ Failed to create placeholder for {file_id}: {e}")
                # Continue with other files
                continue

        logger.info(f"   ✅ Created {len(created_docs)} placeholder documents")
        return created_docs

    except Exception as e:
        logger.error(f"❌ Failed to create placeholder documents: {e}")
        raise GoogleDriveSyncError(f"Failed to create placeholder documents: {e}")


def sync_files(
    user_id: str,
    file_ids: List[str],
    placeholder_doc_ids: Optional[List[str]] = None
) -> Dict:
    """
    Sync specific files from Google Drive by their file IDs.

    Args:
        user_id: UUID of the user
        file_ids: List of Google Drive file IDs to sync

    Returns:
        Dict with sync results

    Raises:
        GoogleDriveSyncError: If sync fails
    """
    # Create sync log entry
    sync_log_id = _create_sync_log(user_id, folder_id=None, folder_name=f"{len(file_ids)} files")

    try:
        logger.info(f"\n📄 Starting Google Drive file sync for user {user_id}")
        logger.info(f"   Syncing {len(file_ids)} file(s)")

        # NOTE: Placeholder documents are created in the API endpoint before this background task
        # We don't create them here to avoid duplicates

        # Get valid credentials
        credentials = get_or_refresh_token(user_id)

        # Get user's client_id
        user_result = supabase.table('users').select('client_id').eq('id', user_id).execute()
        if not user_result.data:
            raise GoogleDriveSyncError(f"User {user_id} not found")

        client_id = user_result.data[0].get('client_id')

        # Initialize LlamaIndex Google Drive reader
        logger.info(f"   🔌 Connecting to Google Drive...")
        authorized_user_info = {
            "client_id": credentials.client_id,
            "client_secret": credentials.client_secret,
            "refresh_token": credentials.refresh_token,
            "type": "authorized_user"
        }
        # Configure file extractors for proper .docx text extraction
        file_extractor = {
            ".docx": SimpleDocxReader(),  # Properly extract text from Google Docs exported as .docx
        }
        reader = GoogleDriveReader(
            authorized_user_info=authorized_user_info,
            file_extractor=file_extractor
        )

        # Load documents by file IDs
        logger.info(f"   📥 Fetching documents...")
        documents = []
        try:
            with timeout(600):  # 10 minute timeout
                documents = reader.load_data(file_ids=file_ids)
        except TimeoutError as e:
            logger.error(f"   ⏱️  Document fetch timed out: {e}")
            raise GoogleDriveSyncError("Document fetch timed out after 10 minutes.")
        except Exception as e:
            logger.warning(f"   ⚠️  Error during document fetch: {e}")
            raise GoogleDriveSyncError(f"Failed to fetch documents: {e}")

        logger.info(f"   ✓ Found {len(documents)} documents")

        # Optimize: Fetch all existing documents once
        logger.info(f"   📋 Fetching existing documents...")
        existing_docs_result = supabase.table('documents')\
            .select('*')\
            .eq('user_id', user_id)\
            .not_.is_('google_drive_file_id', 'null')\
            .execute()

        # Create lookup map by google_drive_file_id
        existing_docs_map = {
            doc['google_drive_file_id']: doc
            for doc in existing_docs_result.data
            if doc.get('google_drive_file_id')
        }
        logger.info(f"   ✓ Found {len(existing_docs_map)} existing Google Drive documents")

        # Track sync statistics
        stats = {
            'documents_added': 0,
            'documents_updated': 0,
            'documents_skipped': 0,
            'errors': []
        }

        # Process each document
        # Zip file_ids with documents to ensure we know which file_id corresponds to which document
        # GoogleDriveReader doesn't reliably include file_id in metadata
        for idx, (file_id, doc) in enumerate(zip(file_ids, documents), 1):
            try:
                logger.info(f"\n   📄 [{idx}/{len(documents)}] Processing: {doc.metadata.get('file_name', 'Unknown')}")

                # Extract metadata
                file_name = (doc.metadata.get('file_name') or
                            doc.metadata.get('name') or
                            doc.metadata.get('title') or
                            'untitled')
                web_view_link = doc.metadata.get('web_view_link', '')
                mime_type = doc.metadata.get('mime_type', 'text/plain')

                # Check if document already exists (need this early to mark failed docs)
                existing_doc = existing_docs_map.get(file_id)

                # Skip folders
                if mime_type == 'application/vnd.google-apps.folder':
                    logger.info(f"      📁 Skipping folder: {file_name}")
                    stats['documents_skipped'] += 1
                    continue

                # Validate document text content
                if not doc.text or len(doc.text.strip()) == 0:
                    logger.warning(f"      ⚠️  Document has no text content, skipping...")
                    # Mark placeholder as failed
                    if existing_doc:
                        supabase.table('documents').update({
                            'processed': True,
                            'processing_status': 'failed',
                            'processing_error': 'No text content available'
                        }).eq('id', existing_doc['id']).execute()
                    stats['documents_skipped'] += 1
                    stats['errors'].append({
                        'file_name': file_name,
                        'file_id': file_id,
                        'error': 'No text content available'
                    })
                    continue

                # Check if text is actually readable (not binary garbage)
                if not is_valid_text_content(doc.text):
                    logger.warning(f"      ⚠️  Document text appears corrupted or binary, skipping...")
                    # Mark placeholder as failed
                    if existing_doc:
                        supabase.table('documents').update({
                            'processed': True,
                            'processing_status': 'failed',
                            'processing_error': 'Text content appears corrupted or binary (possible export failure)'
                        }).eq('id', existing_doc['id']).execute()
                    stats['documents_skipped'] += 1
                    stats['errors'].append({
                        'file_name': file_name,
                        'file_id': file_id,
                        'error': 'Text content appears corrupted or binary (possible export failure)'
                    })
                    continue

                file_content = doc.text.encode('utf-8')

                # Check if it's a placeholder document (created immediately for UI)
                is_placeholder = existing_doc and existing_doc.get('storage_url', '').startswith('placeholder/')

                if existing_doc and not is_placeholder:
                    # Update existing real document (re-sync)
                    logger.info(f"      📝 Document exists, updating...")
                    updated_doc = update_drive_document(
                        document_id=existing_doc['id'],
                        file_content=file_content,
                        storage_url=existing_doc['storage_url']
                    )

                    # Re-process document for embeddings
                    logger.info(f"      🔄 Re-processing for embeddings...")
                    process_document(existing_doc['id'])

                    stats['documents_updated'] += 1

                elif is_placeholder:
                    # Populate placeholder document with actual content
                    logger.info(f"      📋 Populating placeholder document...")

                    # Upload content to storage
                    storage_path = f"{user_id}/{file_id}_{file_name}"
                    logger.info(f"      📤 Uploading to storage: {storage_path}")

                    supabase.storage.from_('documents').upload(
                        path=storage_path,
                        file=file_content,
                        file_options={"upsert": "true"}
                    )

                    storage_url = f"{SUPABASE_URL}/storage/v1/object/public/documents/{storage_path}"

                    # Update placeholder with real content
                    now = datetime.now(timezone.utc).isoformat()
                    update_data = {
                        'storage_url': storage_url,
                        'last_synced_at': now,
                        'processed': False  # Will be set to True after embedding
                    }

                    supabase.table('documents')\
                        .update(update_data)\
                        .eq('id', existing_doc['id'])\
                        .execute()

                    logger.info(f"      ✓ Populated placeholder document")

                    # Process document for embeddings
                    logger.info(f"      🔄 Processing for embeddings...")
                    try:
                        process_document(existing_doc['id'])
                        logger.info(f"      ✓ Processing complete")
                    except Exception as process_error:
                        logger.error(f"      ❌ Processing failed: {process_error}")
                        raise

                    stats['documents_added'] += 1

                else:
                    # Create completely new document
                    logger.info(f"      ➕ New document, creating...")
                    new_doc = create_drive_document(
                        user_id=user_id,
                        client_id=client_id,
                        filename=file_name,
                        file_content=file_content,
                        google_drive_file_id=file_id,
                        external_url=web_view_link,
                        mime_type=mime_type
                    )

                    logger.info(f"      ✓ Created document: {new_doc['id']}")

                    # Process document for embeddings
                    logger.info(f"      🔄 Processing for embeddings...")
                    try:
                        process_document(new_doc['id'])
                        logger.info(f"      ✓ Processing complete")
                    except Exception as process_error:
                        logger.error(f"      ❌ Processing failed: {process_error}")
                        raise

                    stats['documents_added'] += 1

                logger.info(f"      ✅ Done")

            except Exception as e:
                error_msg = str(e)
                logger.error(f"      ❌ Error processing document: {error_msg}")

                # Mark document as processed with error status
                if existing_doc:
                    try:
                        supabase.table('documents').update({
                            'processed': True,
                            'processing_status': 'failed',
                            'processing_error': error_msg[:500]  # Truncate long errors
                        }).eq('id', existing_doc['id']).execute()
                        logger.info(f"      ✓ Marked document as failed")
                    except Exception as update_error:
                        logger.error(f"      ❌ Failed to update document status: {update_error}")

                stats['errors'].append({
                    'file_name': doc.metadata.get('file_name', 'Unknown'),
                    'file_id': doc.metadata.get('file_id', 'Unknown'),
                    'error': error_msg
                })
                stats['documents_skipped'] += 1

        # Update sync log with results
        _complete_sync_log(sync_log_id, 'completed', stats)

        logger.info(f"\n✅ Sync complete!")
        logger.info(f"   Added: {stats['documents_added']}")
        logger.info(f"   Updated: {stats['documents_updated']}")
        logger.info(f"   Skipped: {stats['documents_skipped']}")

        if stats['errors']:
            logger.warning(f"   Errors: {len(stats['errors'])}")
            for error in stats['errors'][:5]:  # Log first 5 errors
                logger.warning(f"      - {error['file_name']}: {error['error']}")

        return stats

    except Exception as e:
        logger.error(f"❌ Sync failed: {str(e)}")
        _complete_sync_log(sync_log_id, 'failed', {'error': str(e)})
        raise


def sync_folder(
    user_id: str,
    folder_id: Optional[str] = None,
    folder_name: Optional[str] = None
) -> Dict:
    """
    Sync documents from a Google Drive folder.

    Args:
        user_id: UUID of the user
        folder_id: Google Drive folder ID (None = entire Drive)
        folder_name: Human-readable folder name for logging

    Returns:
        Dict with sync results

    Raises:
        GoogleDriveSyncError: If sync fails
    """
    # Create sync log entry
    sync_log_id = _create_sync_log(user_id, folder_id, folder_name)

    try:
        logger.info(f"\n📁 Starting Google Drive sync for user {user_id}")
        if folder_id:
            logger.info(f"   Folder: {folder_name or folder_id}")
        else:
            logger.info(f"   Syncing entire Drive")

        # Get valid credentials
        credentials = get_or_refresh_token(user_id)

        # Get user's client_id
        user_result = supabase.table('users').select('client_id').eq('id', user_id).execute()
        if not user_result.data:
            raise GoogleDriveSyncError(f"User {user_id} not found")

        client_id = user_result.data[0].get('client_id')

        # Initialize LlamaIndex Google Drive reader
        # Convert Credentials object to authorized_user_info dict format
        logger.info(f"   🔌 Connecting to Google Drive...")
        authorized_user_info = {
            "client_id": credentials.client_id,
            "client_secret": credentials.client_secret,
            "refresh_token": credentials.refresh_token,
            "type": "authorized_user"
        }
        # Configure file extractors for proper .docx text extraction
        file_extractor = {
            ".docx": SimpleDocxReader(),  # Properly extract text from Google Docs exported as .docx
        }
        reader = GoogleDriveReader(
            authorized_user_info=authorized_user_info,
            file_extractor=file_extractor
        )

        # Load documents with timeout protection (10 minutes max)
        logger.info(f"   📥 Fetching documents...")
        documents = []
        try:
            with timeout(600):  # 10 minute timeout
                if folder_id:
                    # Fetch from specific folder
                    documents = reader.load_data(folder_id=folder_id)
                else:
                    # Fetch all accessible documents (might be a lot!)
                    # For safety, we'll fetch from root folder
                    documents = reader.load_data(folder_id="root")
        except TimeoutError as e:
            logger.error(f"   ⏱️  Document fetch timed out: {e}")
            raise GoogleDriveSyncError("Document fetch timed out after 10 minutes. Please try syncing a smaller folder.")
        except Exception as e:
            # Log the error but try to continue with any documents we did get
            logger.warning(f"   ⚠️  Error during document fetch (will process any retrieved documents): {e}")
            # If we got some documents before the error, we'll process those
            if not documents:
                raise GoogleDriveSyncError(f"Failed to fetch documents: {e}")

        logger.info(f"   ✓ Found {len(documents)} documents")

        # Log document metadata for debugging
        if documents:
            logger.info(f"   📋 Sample document metadata:")
            sample_doc = documents[0]
            logger.info(f"      Available metadata keys: {list(sample_doc.metadata.keys())}")
            logger.info(f"      file_name: {sample_doc.metadata.get('file_name')}")
            logger.info(f"      name: {sample_doc.metadata.get('name')}")
            logger.info(f"      title: {sample_doc.metadata.get('title')}")
            logger.info(f"      file_id: {sample_doc.metadata.get('file_id')}")
            logger.info(f"      mime_type: {sample_doc.metadata.get('mime_type')}")
            logger.info(f"      text length: {len(sample_doc.text) if sample_doc.text else 0}")

        # Optimize: Fetch all existing documents once to avoid N+1 query problem
        logger.info(f"   📋 Fetching existing documents...")
        existing_docs_result = supabase.table('documents')\
            .select('*')\
            .eq('user_id', user_id)\
            .not_.is_('google_drive_file_id', 'null')\
            .execute()

        # Create lookup map by google_drive_file_id
        existing_docs_map = {
            doc['google_drive_file_id']: doc
            for doc in existing_docs_result.data
            if doc.get('google_drive_file_id')
        }
        logger.info(f"   ✓ Found {len(existing_docs_map)} existing Google Drive documents")

        # Track sync statistics
        stats = {
            'documents_added': 0,
            'documents_updated': 0,
            'documents_skipped': 0,
            'errors': []
        }

        # Process each document
        for idx, doc in enumerate(documents, 1):
            try:
                logger.info(f"\n   📄 [{idx}/{len(documents)}] Processing: {doc.metadata.get('file_name', 'Unknown')}")

                # Extract metadata
                file_id = doc.metadata.get('file_id')
                # Try multiple possible keys for filename
                file_name = (doc.metadata.get('file_name') or
                            doc.metadata.get('name') or
                            doc.metadata.get('title') or
                            'untitled')
                web_view_link = doc.metadata.get('web_view_link', '')
                mime_type = doc.metadata.get('mime_type', 'text/plain')

                # Skip folders - we only want files
                if mime_type == 'application/vnd.google-apps.folder':
                    logger.info(f"      📁 Skipping folder: {file_name}")
                    stats['documents_skipped'] += 1
                    continue

                # Get document text content
                # Validate document text content
                if not doc.text or len(doc.text.strip()) == 0:
                    logger.warning(f"      ⚠️  Document has no text content, skipping...")
                    stats['documents_skipped'] += 1
                    stats['errors'].append({
                        'file_name': file_name,
                        'file_id': file_id,
                        'error': 'No text content available'
                    })
                    continue

                # Check if text is actually readable (not binary garbage)
                if not is_valid_text_content(doc.text):
                    logger.warning(f"      ⚠️  Document text appears corrupted or binary, skipping...")
                    stats['documents_skipped'] += 1
                    stats['errors'].append({
                        'file_name': file_name,
                        'file_id': file_id,
                        'error': 'Text content appears corrupted or binary (possible export failure)'
                    })
                    continue

                file_content = doc.text.encode('utf-8')

                # Check if document already exists (using pre-fetched map)
                existing_doc = existing_docs_map.get(file_id)

                if existing_doc:
                    # Update existing document
                    logger.info(f"      📝 Document exists, updating...")
                    updated_doc = update_drive_document(
                        document_id=existing_doc['id'],
                        file_content=file_content,
                        storage_url=existing_doc['storage_url']
                    )

                    # Re-process document for embeddings
                    logger.info(f"      🔄 Re-processing for embeddings...")
                    process_document(existing_doc['id'])

                    stats['documents_updated'] += 1

                else:
                    # Create new document
                    logger.info(f"      ➕ New document, creating...")
                    logger.info(f"      📝 user_id={user_id}, client_id={client_id}")
                    new_doc = create_drive_document(
                        user_id=user_id,
                        client_id=client_id,
                        filename=file_name,
                        file_content=file_content,
                        google_drive_file_id=file_id,
                        external_url=web_view_link,
                        mime_type=mime_type
                    )

                    logger.info(f"      ✓ Created document: {new_doc['id']}")
                    logger.info(f"      📋 Document uploaded_by={new_doc.get('uploaded_by')}, client_id={new_doc.get('client_id')}")

                    # Process document for embeddings
                    logger.info(f"      🔄 Processing for embeddings...")
                    try:
                        process_document(new_doc['id'])
                        logger.info(f"      ✓ Processing complete")
                    except Exception as process_error:
                        logger.error(f"      ❌ Processing failed: {process_error}")
                        raise

                    stats['documents_added'] += 1

                logger.info(f"      ✅ Done")

            except Exception as e:
                # Enhanced error logging with more details
                error_msg = str(e)
                logger.error(f"      ❌ Error processing document: {error_msg}")

                # Check for specific Google Drive API errors
                if 'exportSizeLimitExceeded' in error_msg or 'too large to be exported' in error_msg:
                    logger.warning(f"      ⚠️  File too large to export (Google Drive limit)")
                elif 'fileNotDownloadable' in error_msg or 'Use Export with Docs Editors files' in error_msg:
                    logger.warning(f"      ⚠️  File type not downloadable (Google Docs/Sheets/Slides export issue)")

                stats['errors'].append({
                    'file_name': doc.metadata.get('file_name', 'Unknown'),
                    'file_id': doc.metadata.get('file_id', 'Unknown'),
                    'error': error_msg
                })
                stats['documents_skipped'] += 1

        # Check if we got any valid documents
        if stats['documents_added'] == 0 and stats['documents_updated'] == 0:
            if folder_id:
                error_msg = (
                    f"No documents found in the specified folder. "
                    f"Please ensure: (1) The folder contains supported file types (PDF, TXT, DOCX, etc.), "
                    f"(2) Files are shared with your Google account, "
                    f"(3) You're using a folder ID, not a file ID. "
                    f"Skipped: {stats['documents_skipped']}, Errors: {len(stats['errors'])}"
                )
                logger.warning(f"   ⚠️  {error_msg}")
                raise GoogleDriveSyncError(error_msg)

        # Update sync log with results
        _complete_sync_log(sync_log_id, 'completed', stats)

        logger.info(f"\n✅ Sync complete!")
        logger.info(f"   Added: {stats['documents_added']}")
        logger.info(f"   Updated: {stats['documents_updated']}")
        logger.info(f"   Skipped: {stats['documents_skipped']}")

        # Log details about errors if any occurred
        if stats['errors']:
            logger.info(f"\n   ⚠️  {len(stats['errors'])} file(s) had errors:")
            for idx, error in enumerate(stats['errors'][:5], 1):  # Show first 5 errors
                logger.info(f"      {idx}. {error.get('file_name', 'Unknown')}: {error.get('error', 'Unknown error')[:100]}")
            if len(stats['errors']) > 5:
                logger.info(f"      ... and {len(stats['errors']) - 5} more")

        return {
            'status': 'success',
            'sync_log_id': sync_log_id,
            **stats
        }

    except Exception as e:
        # Mark sync as failed
        error_message = str(e)
        _complete_sync_log(sync_log_id, 'failed', error_message=error_message)

        logger.error(f"\n❌ Sync failed: {e}")
        raise GoogleDriveSyncError(f"Sync failed: {e}")


def _create_sync_log(user_id: str, folder_id: Optional[str], folder_name: Optional[str]) -> str:
    """Create a sync log entry and return its ID"""
    log_data = {
        'user_id': user_id,
        'folder_id': folder_id,
        'folder_name': folder_name or 'Entire Drive',
        'sync_type': 'full',
        'status': 'running',
        'started_at': datetime.now(timezone.utc).isoformat()
    }

    result = supabase.table('google_drive_sync_log').insert(log_data).execute()
    return result.data[0]['id']


def _complete_sync_log(
    sync_log_id: str,
    status: str,
    stats: Optional[Dict] = None,
    error_message: Optional[str] = None
):
    """Update sync log entry with completion status"""
    now = datetime.now(timezone.utc).isoformat()

    update_data = {
        'status': status,
        'completed_at': now
    }

    if stats:
        update_data.update({
            'documents_added': stats.get('documents_added', 0),
            'documents_updated': stats.get('documents_updated', 0),
            'documents_skipped': stats.get('documents_skipped', 0)
        })

        # If there were errors, create a summary error message
        if stats.get('errors'):
            error_count = len(stats['errors'])
            error_summary = f"{error_count} file(s) had errors. "
            if error_count > 0:
                first_errors = stats['errors'][:3]
                error_details = '; '.join([
                    f"{e.get('file_name', 'Unknown')}: {str(e.get('error', 'Unknown'))[:50]}"
                    for e in first_errors
                ])
                error_summary += error_details
                if error_count > 3:
                    error_summary += f"; ... and {error_count - 3} more"
            update_data['error_message'] = error_summary

    if error_message:
        update_data['error_message'] = error_message

    supabase.table('google_drive_sync_log')\
        .update(update_data)\
        .eq('id', sync_log_id)\
        .execute()


# ============================================================================
# Connection Management
# ============================================================================

def is_user_connected(user_id: str) -> bool:
    """Check if user has an active Google Drive connection"""
    token_record = get_user_tokens(user_id)
    return token_record is not None


def get_connection_status(user_id: str) -> Dict:
    """
    Get Google Drive connection status for a user.

    Returns:
        Dict with connection status, last sync time, document count
    """
    try:
        connected = is_user_connected(user_id)

        if not connected:
            return {
                'connected': False,
                'last_sync': None,
                'document_count': 0
            }

        # Get last sync time
        sync_result = supabase.table('google_drive_sync_log')\
            .select('completed_at')\
            .eq('user_id', user_id)\
            .eq('status', 'completed')\
            .order('completed_at', desc=True)\
            .limit(1)\
            .execute()

        last_sync = sync_result.data[0]['completed_at'] if sync_result.data else None

        # Get document count
        doc_result = supabase.table('documents')\
            .select('id', count='exact')\
            .eq('user_id', user_id)\
            .eq('source_platform', 'google_drive')\
            .execute()

        document_count = doc_result.count or 0

        return {
            'connected': True,
            'last_sync': last_sync,
            'document_count': document_count
        }

    except Exception as e:
        logger.warning(f"⚠️  Error getting connection status: {e}")
        return {
            'connected': False,
            'last_sync': None,
            'document_count': 0,
            'error': str(e)
        }


def disconnect_user(user_id: str) -> Dict:
    """
    Disconnect user's Google Drive (delete tokens).
    Note: Does not delete synced documents.
    """
    try:
        # Delete the token record completely to ensure clean reconnection
        supabase.table('google_drive_tokens')\
            .delete()\
            .eq('user_id', user_id)\
            .execute()

        return {'status': 'success', 'message': 'Google Drive disconnected'}

    except Exception as e:
        raise GoogleDriveSyncError(f"Failed to disconnect: {e}")
